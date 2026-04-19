"""Convert WRITTEN_REPORT.md to a Word document with the required formatting:
- 11pt sans serif font
- 1.15 line spacing
- Proper headers, tables, and lists
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_LINE_SPACING

BASE = Path(__file__).resolve().parent.parent
md_path = BASE / 'WRITTEN_REPORT.md'
docx_path = BASE / 'WRITTEN_REPORT.docx'

doc = Document()

# Set default style: Calibri 11pt, 1.15 spacing
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_heading(text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)


def parse_inline(text):
    """Parse bold, italic, code in inline text. Returns list of (text, bold, italic, code) tuples."""
    parts = []
    # Simple regex to find **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|_[^_]+_)')
    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False, False, False))
        tok = match.group(0)
        if tok.startswith('**'):
            parts.append((tok[2:-2], True, False, False))
        elif tok.startswith('`'):
            parts.append((tok[1:-1], False, False, True))
        elif tok.startswith('*') or tok.startswith('_'):
            parts.append((tok[1:-1], False, True, False))
        last_end = match.end()
    if last_end < len(text):
        parts.append((text[last_end:], False, False, False))
    return parts


def add_formatted_paragraph(text, style_name='Normal'):
    para = doc.add_paragraph(style=style_name)
    parts = parse_inline(text)
    for t, bold, italic, code in parts:
        if not t:
            continue
        run = para.add_run(t)
        run.font.name = 'Consolas' if code else 'Calibri'
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.color.rgb = RGBColor(0x80, 0x40, 0x20)
    return para


def add_bullet(text):
    para = doc.add_paragraph(style='List Bullet')
    parts = parse_inline(text)
    for t, bold, italic, code in parts:
        if not t:
            continue
        run = para.add_run(t)
        run.font.name = 'Consolas' if code else 'Calibri'
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic


def add_table(rows):
    """rows is a list of list of cell text (first row is header)."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ''  # clear default paragraph
            para = cell.paragraphs[0]
            parts = parse_inline(cell_text)
            for t, bold, italic, code in parts:
                if not t:
                    continue
                run = para.add_run(t)
                run.font.name = 'Consolas' if code else 'Calibri'
                run.font.size = Pt(10)
                run.bold = bold or (i == 0)  # header row bold
                run.italic = italic


# --- Parse markdown ---
lines = md_path.read_text().splitlines()
i = 0
while i < len(lines):
    line = lines[i].rstrip()

    if not line.strip():
        i += 1
        continue

    # Horizontal rule — skip
    if re.match(r'^-{3,}$', line):
        i += 1
        continue

    # Headings
    m = re.match(r'^(#{1,6})\s+(.+)$', line)
    if m:
        level = len(m.group(1))
        add_heading(m.group(2).strip(), level)
        i += 1
        continue

    # Tables: line starts with | and next line is | --- | --- |
    if line.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s\-:\|]+\|$', lines[i + 1]):
        # Collect table rows
        rows = []
        header = [c.strip() for c in line.strip('|').split('|')]
        rows.append(header)
        i += 2  # skip header + separator
        while i < len(lines) and lines[i].startswith('|'):
            cells = [c.strip() for c in lines[i].strip('|').split('|')]
            rows.append(cells)
            i += 1
        add_table(rows)
        continue

    # Unordered list
    if re.match(r'^\s*[-*]\s+', line):
        text = re.sub(r'^\s*[-*]\s+', '', line)
        add_bullet(text)
        i += 1
        continue

    # Ordered list
    if re.match(r'^\s*\d+\.\s+', line):
        text = re.sub(r'^\s*\d+\.\s+', '', line)
        para = doc.add_paragraph(style='List Number')
        parts = parse_inline(text)
        for t, bold, italic, code in parts:
            run = para.add_run(t)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.bold = bold
            run.italic = italic
        i += 1
        continue

    # Block quote — treat as italic paragraph
    if line.startswith('>'):
        text = re.sub(r'^>\s?', '', line)
        para = add_formatted_paragraph(text)
        for run in para.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        i += 1
        continue

    # Regular paragraph
    add_formatted_paragraph(line)
    i += 1

doc.save(docx_path)
print(f"Saved: {docx_path}")
print(f"Size: {docx_path.stat().st_size} bytes")
